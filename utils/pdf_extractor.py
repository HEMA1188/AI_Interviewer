# pdf_extractor.py

import os
import fitz # PyMuPDF
import logging
from docx import Document # python-docx
import asyncio

logger = logging.getLogger(__name__)

# REMOVED: pdf_path="HEMA_AI_Developer.pdf" (global variable)

async def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extracts text from a PDF file using PyMuPDF (Fitz).
    Returns an empty string if the file is not found or an error occurs.
    """
    if not os.path.exists(pdf_path):
        logger.error(f"PDF file not found: {pdf_path}")
        return ""
    
    text = ""
    try:
        # Using 'with' statement to ensure the document is properly closed
        with fitz.open(pdf_path) as doc:
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                text += page.get_text() # get_text() typically handles line breaks within a page
        logger.info(f"Successfully extracted text from PDF: {pdf_path}")
    except Exception as e:
        logger.error(f"Error extracting text from PDF '{pdf_path}': {e}")
    return text

async def extract_text_from_docx(docx_path: str) -> str:
    """
    Extracts text from a DOCX file using python-docx.
    Returns an empty string if the file is not found or an error occurs.
    """
    if not os.path.exists(docx_path):
        logger.error(f"DOCX file not found: {docx_path}")
        return ""

    text = ""
    try:
        document = Document(docx_path)
        for para in document.paragraphs:
            text += para.text + "\n" # Adding newline for paragraph separation
        logger.info(f"Successfully extracted text from DOCX: {docx_path}")
    except Exception as e:
        logger.error(f"Error extracting text from DOCX '{docx_path}': {e}")
    return text

async def get_resume_text(resume_path: str) -> str:
    """
    Determines file type and extracts text from a resume file (PDF, DOCX, or TXT).
    Returns the extracted text or an empty string on failure.
    """
    if not os.path.exists(resume_path):
        logger.error(f"Resume file not found at: {resume_path}")
        return ""

    file_extension = os.path.splitext(resume_path)[1].lower()

    if file_extension == '.pdf':
        return await extract_text_from_pdf(resume_path)
    elif file_extension == '.docx':
        return await extract_text_from_docx(resume_path)
    elif file_extension == '.txt':
        try:
            with open(resume_path, 'r', encoding='utf-8') as f:
                text = f.read()
            logger.info(f"Successfully read text from TXT: {resume_path}")
            return text
        except Exception as e:
            logger.error(f"Error reading text from TXT file '{resume_path}': {e}")
            return ""
    else:
        logger.warning(f"Unsupported resume file format: {file_extension}. Only .pdf, .docx, and .txt are supported.")
        return ""

if __name__ == "__main__":
    async def test_extractor():
        # Setup basic logging for standalone execution
        logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        
        # Create dummy files for testing
        dummy_txt_path = "dummy_resume.txt"
        with open(dummy_txt_path, "w", encoding="utf-8") as f:
            f.write("This is a dummy text resume.\nIt has two lines.")
        
        dummy_docx_path = "dummy_resume.docx"
        from docx import Document
        doc = Document()
        doc.add_paragraph("This is a dummy Word resume.")
        doc.add_paragraph("It also has multiple lines.")
        doc.save(dummy_docx_path)

        # Placeholder for a real PDF path. Make sure to replace this with an actual PDF for testing.
        real_pdf_path = "HEMA_AI_Developer.pdf" # Replace with your actual PDF path for thorough testing

        print("\n--- Testing TXT ---")
        txt_text = await get_resume_text(dummy_txt_path)
        print(f"Extracted from {dummy_txt_path}:\n{txt_text}")

        print("\n--- Testing DOCX ---")
        docx_text = await get_resume_text(dummy_docx_path)
        print(f"Extracted from {dummy_docx_path}:\n{docx_text}")

        print("\n--- Testing PDF (using a real file if available) ---")
        if os.path.exists(real_pdf_path):
            pdf_text_real = await get_resume_text(real_pdf_path)
            print(f"Extracted from {real_pdf_path}:\n{pdf_text_real[:500]}...") # Print first 500 chars
        else:
            print(f"Skipping real PDF test: '{real_pdf_path}' not found. Please place a PDF file there for full testing.")
            pdf_text_non_existent = await get_resume_text("non_existent.pdf")
            print(f"Extracted from non_existent.pdf: '{pdf_text_non_existent}' (Expected: '')")


        # Cleanup dummy files
        os.remove(dummy_txt_path)
        os.remove(dummy_docx_path)
        # Uncomment the line below if you create a dummy PDF in the test section
        # if os.path.exists("dummy_resume.pdf"): os.remove("dummy_resume.pdf")

    asyncio.run(test_extractor())